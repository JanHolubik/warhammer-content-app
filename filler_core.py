#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import re
import csv
import unicodedata
from pathlib import Path
from typing import Dict, List, Optional

import pandas as pd
from docx import Document


# =========================
# TEMPLATE MAPA
# =========================

def pick_existing(base_dir: Path, candidates: List[str]) -> Path:
    for name in candidates:
        p = base_dir / name
        if p.exists():
            return p
    raise FileNotFoundError(f"Nenašel jsem žádnou z variant: {candidates} v {base_dir}")


def get_template_paths(template_dir: Path, template_type: str) -> Dict[str, Path]:
    template_type = template_type.strip().lower()

    short_map = {
        "miniatures": ["kratky.docx"],
        "books": ["kratky_univ.docx"],
        "warscroll": ["kratky_univ.docx"],
        "dice": ["kratky.docx"],
        "upgrades": ["kratky.docx"],
    }

    detail_map = {
        "miniatures": ["detailni.docx"],
        "books": ["detailni_kniha.docx"],
        "warscroll": ["detailni_warscoll.docx", "detailni_warscroll.docx"],
        "dice": ["detailni_kostky.docx"],
        "upgrades": ["detailni_upgrades.docx"],
    }

    short_map_en = {
        "miniatures": ["kratky_en.docx"],
        "books": ["kratky_univ_en.docx"],
        "warscroll": ["kratky_univ_en.docx"],
        "dice": ["kratky_en.docx"],
        "upgrades": ["kratky_en.docx"],
    }

    detail_map_en = {
        "miniatures": ["detailni_en.docx"],
        "books": ["detailni_kniha_en.docx"],
        "warscroll": ["detailni_warscoll_en.docx", "detailni_warscroll_en.docx"],
        "dice": ["detailni_kostky_en.docx"],
        "upgrades": ["detailni_upgrades_en.docx"],
    }

    short_map_sk = {
        "miniatures": ["kratky_sk.docx"],
        "books": ["kratky_univ_sk.docx"],
        "warscroll": ["kratky_univ_sk.docx"],
        "dice": ["kratky_sk.docx"],
        "upgrades": ["kratky_sk.docx"],
    }

    detail_map_sk = {
        "miniatures": ["detailni_sk.docx"],
        "books": ["detailni_kniha_sk.docx"],
        "warscroll": ["detailni_warscoll_sk.docx", "detailni_warscroll_sk.docx"],
        "dice": ["detailni_kostky_sk.docx"],
        "upgrades": ["detailni_upgrades_sk.docx"],
    }

    if template_type not in short_map:
        raise ValueError(f"Neznámý TEMPLATE_TYPE: {template_type}")

    return {
        "short_cs": pick_existing(template_dir, short_map[template_type]),
        "detail_cs": pick_existing(template_dir, detail_map[template_type]),
        "short_en": pick_existing(template_dir, short_map_en[template_type]),
        "detail_en": pick_existing(template_dir, detail_map_en[template_type]),
        "short_sk": pick_existing(template_dir, short_map_sk[template_type]),
        "detail_sk": pick_existing(template_dir, detail_map_sk[template_type]),
    }


# =========================
# DOCX
# =========================

def normalize_docx_text(text: str) -> str:
    return (
        str(text)
        .replace("\ufeff", "")
        .replace("\u200b", "")
        .replace("\xa0", " ")
        .replace("\u202f", " ")
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )


def read_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: List[str] = []

    for p in doc.paragraphs:
        if p.text is not None:
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text is not None:
                    parts.append(cell.text)

    return normalize_docx_text("\n".join(parts))


# =========================
# KLÍČE / PLACEHOLDERY
# =========================

def canonical_key(text: str) -> str:
    text = normalize_docx_text(text).strip().lower()
    text = "".join(
        ch for ch in unicodedata.normalize("NFKD", text)
        if not unicodedata.combining(ch)
    )
    text = re.sub(r"[^a-z0-9_]", "", text)
    return text


def extract_placeholders(template_text: str) -> List[str]:
    found = re.findall(
        r"\{([A-Za-z0-9_áčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ]+)\}",
        template_text
    )

    out: List[str] = []
    seen = set()

    for item in found:
        item = item.strip()
        if item and item not in seen:
            seen.add(item)
            out.append(item)

    return out


# =========================
# PARSE PROMPT OUTPUT
# =========================

def parse_key_value_block(text: str) -> Dict[str, str]:
    text = normalize_docx_text(text).strip()
    lines = text.split("\n")

    result: Dict[str, str] = {}
    current_key: Optional[str] = None
    current_value_lines: List[str] = []

    for raw_line in lines:
        line = normalize_docx_text(raw_line).strip()

        if not line:
            if current_key is not None:
                current_value_lines.append("")
            continue

        if ":" in line:
            possible_key, possible_value = line.split(":", 1)
            possible_key_clean = canonical_key(possible_key)

            if possible_key_clean:
                if current_key is not None:
                    result[current_key] = normalize_docx_text("\n".join(current_value_lines).strip())

                current_key = possible_key_clean
                current_value_lines = [possible_value.strip()] if possible_value.strip() else []
                continue

        if current_key is not None:
            current_value_lines.append(line)

    if current_key is not None:
        result[current_key] = normalize_docx_text("\n".join(current_value_lines).strip())

    return result


def parse_prompt_output_by_lang(text: str) -> Dict[str, Dict[str, str]]:
    text = normalize_docx_text(text).strip()

    pattern = re.compile(r"(?ms)^\[LANG=(cs|en|sk)\]\s*$")
    matches = list(pattern.finditer(text))

    if not matches:
        single = parse_key_value_block(text)
        return {
            "cs": dict(single),
            "en": dict(single),
            "sk": dict(single),
        }

    sections: Dict[str, Dict[str, str]] = {}

    for i, match in enumerate(matches):
        lang = match.group(1).strip().lower()
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block_text = text[start:end].strip()
        sections[lang] = parse_key_value_block(block_text)

    for required in ("cs", "en", "sk"):
        if required not in sections:
            raise ValueError(f"Chybí jazykový blok [LANG={required}] ve vystup_prompt.docx.")

    return sections


# =========================
# ALIASY
# =========================

def apply_aliases(values: Dict[str, str], template_type: str) -> Dict[str, str]:
    out = dict(values)

    common_aliases = {
        "nazev_sady": "název_sady",
        "potrebne_vybaveni": "potřebné_vybavení",
        "doporuceni_pro_koho": "doporučení_pro_koho",
    }

    for src, dst in common_aliases.items():
        if src in out and dst not in out:
            out[dst] = out[src]

    if template_type == "miniatures":
        mini_aliases = {
            "nazev_produktu": "product_title",
            "uvodni_text": "intro_text",

            "sekce_o_jednotce_nadpis": "section_inside_title",
            "sekce_o_jednotce_text": "section_inside_text",

            "sekce_zarazeni_nadpis": "section_position_title",
            "sekce_zarazeni_text": "section_position_text",

            "sekce_vyhody_nadpis": "section_benefits_title",
            "vyhoda_1": "benefit_1",
            "vyhoda_2": "benefit_2",
            "vyhoda_3": "benefit_3",

            "sekce_matchupy_nadpis": "section_matchup_title",
            "frakce1_nazev": "faction1_name",
            "frakce1_popis": "faction1_description",
            "frakce2_nazev": "faction2_name",
            "frakce2_popis": "faction2_description",
            "souhrn_matchupu_nadpis": "matchup_summary_title",
            "souhrn_matchupu_text": "matchup_summary_text",

            "sekce_obsah_nadpis": "section_contents_title",
            "obsah_radek1_label": "contents_row1_label",
            "obsah_radek1_text": "contents_row1_text",
            "obsah_radek2_label": "contents_row2_label",
            "obsah_radek2_text": "contents_row2_text",
            "obsah_radek3_label": "contents_row3_label",
            "obsah_radek3_text": "contents_row3_text",
            "obsah_radek4_label": "contents_row4_label",
            "obsah_radek4_text": "contents_row4_text",
            "obsah_radek5_label": "contents_row5_label",
            "obsah_radek5_text": "contents_row5_text",
            "obsah_radek6_label": "contents_row6_label",
            "obsah_radek6_text": "contents_row6_text",

            "sekce_malovani_nadpis": "section_painting_title",

            "primer_nazev": "primer_name",
            "primer_pouziti": "primer_usage",
            "primer_barva": "primer_color",

            "hlavni_barva_1_nazev": "main_color_1_name",
            "hlavni_barva_1_pouziti": "main_color_1_usage",
            "hlavni_barva_1_barva": "main_color_1_color",

            "hlavni_barva_2_nazev": "main_color_2_name",
            "hlavni_barva_2_pouziti": "main_color_2_usage",
            "hlavni_barva_2_barva": "main_color_2_color",

            "hlavni_barva_3_nazev": "main_color_3_name",
            "hlavni_barva_3_pouziti": "main_color_3_usage",
            "hlavni_barva_3_barva": "main_color_3_color",

            "hlavni_barva_4_nazev": "main_color_4_name",
            "hlavni_barva_4_pouziti": "main_color_4_usage",
            "hlavni_barva_4_barva": "main_color_4_color",

            "detailni_barva_1_nazev": "detail_color_1_name",
            "detailni_barva_1_pouziti": "detail_color_1_usage",
            "detailni_barva_1_barva": "detail_color_1_color",

            "detailni_barva_2_nazev": "detail_color_2_name",
            "detailni_barva_2_pouziti": "detail_color_2_usage",
            "detailni_barva_2_barva": "detail_color_2_color",

            "detailni_barva_3_nazev": "detail_color_3_name",
            "detailni_barva_3_pouziti": "detail_color_3_usage",
            "detailni_barva_3_barva": "detail_color_3_color",

            "detailni_barva_4_nazev": "detail_color_4_name",
            "detailni_barva_4_pouziti": "detail_color_4_usage",
            "detailni_barva_4_barva": "detail_color_4_color",

            "wash_1_nazev": "wash_1_name",
            "wash_1_pouziti": "wash_1_usage",
            "wash_1_barva": "wash_1_color",

            "wash_2_nazev": "wash_2_name",
            "wash_2_pouziti": "wash_2_usage",
            "wash_2_barva": "wash_2_color",

            "highlight_1_nazev": "highlight_1_name",
            "highlight_1_pouziti": "highlight_1_usage",
            "highlight_1_barva": "highlight_1_color",

            "highlight_2_nazev": "highlight_2_name",
            "highlight_2_pouziti": "highlight_2_usage",
            "highlight_2_barva": "highlight_2_color",

            "lepidlo_nazev": "glue_name",
            "lepidlo_pouziti": "glue_usage",

            "stetec_1_nazev": "brush_1_name",
            "stetec_1_pouziti": "brush_1_usage",

            "stetec_2_nazev": "brush_2_name",
            "stetec_2_pouziti": "brush_2_usage",

            "sekce_taktika_nadpis": "section_tactics_title",
            "taktika_radek1_label": "tactic_row1_label",
            "taktika_radek1_text": "tactic_row1_text",
            "taktika_radek2_label": "tactic_row2_label",
            "taktika_radek2_text": "tactic_row2_text",
            "taktika_radek3_label": "tactic_row3_label",
            "taktika_radek3_text": "tactic_row3_text",
            "taktika_radek4_label": "tactic_row4_label",
            "taktika_radek4_text": "tactic_row4_text",

            "sekce_faq_nadpis": "section_faq_title",
            "faq_otazka1": "faq_q1",
            "faq_odpoved1": "faq_a1",
            "faq_otazka2": "faq_q2",
            "faq_odpoved2": "faq_a2",
            "faq_otazka3": "faq_q3",
            "faq_odpoved3": "faq_a3",
            "faq_otazka4": "faq_q4",
            "faq_odpoved4": "faq_a4",

            "sekce_pokracovani_nadpis": "section_progress_title",
            "pokracovani_radek1_label": "progress_row1_label",
            "pokracovani_radek1_text": "progress_row1_text",
            "pokracovani_radek2_label": "progress_row2_label",
            "pokracovani_radek2_text": "progress_row2_text",

            "sekce_navaznosti_nadpis": "section_next_title",
            "navaznost_radek1_label": "next_row1_label",
            "navaznost_radek1_text": "next_row1_text",
            "navaznost_radek2_label": "next_row2_label",
            "navaznost_radek2_text": "next_row2_text",
            "navaznost_radek3_label": "next_row3_label",
            "navaznost_radek3_text": "next_row3_text",

            "sekce_video_nadpis": "section_video_title",
        }

        for src, dst in mini_aliases.items():
            if src in out and dst not in out:
                out[dst] = out[src]

        if "section_inside_title" in out and "sekce_o_jednotce_nadpis" not in out:
            out["sekce_o_jednotce_nadpis"] = out["section_inside_title"]

        if "section_inside_text" in out and "sekce_o_jednotce_text" not in out:
            out["sekce_o_jednotce_text"] = out["section_inside_text"]

    return out


# =========================
# HTML
# =========================

def html_escape_basic(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
    )


def is_multiline_text_placeholder(placeholder_name: str) -> bool:
    name = canonical_key(placeholder_name)

    exact_no_br = {
        "video_url",
        "intro_image_src",
        "img1_src",
        "img2_src",
        "img3_src",
        "img4_src",
        "img1_alt",
        "img2_alt",
        "img3_alt",
        "img4_alt",
        "primer_color",
        "main_color_1_color",
        "main_color_2_color",
        "main_color_3_color",
        "main_color_4_color",
        "detail_color_1_color",
        "detail_color_2_color",
        "detail_color_3_color",
        "detail_color_4_color",
        "wash_1_color",
        "wash_2_color",
        "highlight_1_color",
        "highlight_2_color",
        "nazev_sady",
        "nazev_produktu",
        "product_title",
        "link_1_url",
        "link_2_url",
        "link_3_url",
        "link_4_url",
    }

    if name in exact_no_br:
        return False

    suffix_no_br = (
        "_src",
        "_alt",
        "_color",
        "_url",
    )

    return not name.endswith(suffix_no_br)


def prepare_value_for_html(placeholder_name: str, raw_value: str) -> str:
    raw_value = normalize_docx_text(raw_value)
    safe = html_escape_basic(raw_value)

    if is_multiline_text_placeholder(placeholder_name):
        safe = safe.replace("\n", "<br />\n")

    return safe


def flatten_html_for_csv(html: str) -> str:
    html = str(html).replace("\r\n", "\n").replace("\r", "\n")
    html = html.replace("\n", " ")
    html = re.sub(r"\s{2,}", " ", html)
    return html.strip()


# =========================
# DEFAULTY
# =========================

def inject_safe_defaults(values: Dict[str, str], product_name: str) -> Dict[str, str]:
    out = dict(values)

    defaults = {
        "img1_src": "img_src",
        "img2_src": "img_src",
        "img3_src": "img_src",
        "img4_src": "img_src",
        "intro_image_src": "img_src",
        "video_url": "https://www.youtube.com/embed/",
        "link_1_label": "",
        "link_1_url": "",
        "link_2_label": "",
        "link_2_url": "",
        "link_3_label": "",
        "link_3_url": "",
        "link_4_label": "",
        "link_4_url": "",
    }

    for key, value in defaults.items():
        ck = canonical_key(key)
        if not str(out.get(ck, "")).strip():
            out[ck] = value

    for key in ("img1_alt", "img2_alt", "img3_alt", "img4_alt"):
        ck = canonical_key(key)
        if not str(out.get(ck, "")).strip():
            out[ck] = product_name

    return out


# =========================
# TEMPLATE FILL
# =========================

def fill_template(template_text: str, values: Dict[str, str], strict: bool = True) -> str:
    placeholders = extract_placeholders(template_text)

    missing: List[str] = []
    for ph in placeholders:
        ck = canonical_key(ph)
        if ck not in values or not str(values.get(ck, "")).strip():
            missing.append(ph)

    if missing and strict:
        raise ValueError("Chybí placeholdery:\n- " + "\n- ".join(missing))

    filled = template_text

    for ph in placeholders:
        ck = canonical_key(ph)
        raw_value = values.get(ck, "")
        final_value = prepare_value_for_html(ph, raw_value)
        filled = filled.replace("{" + ph + "}", final_value)

    return filled


def print_placeholder_summary(template_name: str, template_placeholders: List[str], values: Dict[str, str]) -> None:
    missing = []

    for ph in template_placeholders:
        ck = canonical_key(ph)
        if ck not in values or not str(values.get(ck, "")).strip():
            missing.append(ph)

    print("=" * 70)
    print(f"KONTROLA: {template_name}")
    print("=" * 70)
    print(f"Počet placeholderů v šabloně: {len(template_placeholders)}")
    print(f"Počet dodaných hodnot: {len(values)}")
    print(f"Chybějící nebo prázdné: {len(missing)}")

    if missing:
        for item in missing:
            print("  -", item)

    print()


# =========================
# CSV
# =========================

def build_csv_row_mask(
    df: pd.DataFrame,
    target_product_name: Optional[str] = None,
    target_ean: Optional[str] = None,
) -> pd.Series:
    if target_ean:
        if "ean" not in df.columns:
            raise ValueError("V CSV chybí sloupec 'ean'.")
        mask = df["ean"].astype(str).str.strip() == target_ean.strip()
        if mask.any():
            return mask
        raise ValueError(f"Produkt s EAN '{target_ean}' nebyl nalezen.")

    if target_product_name:
        name_col = "name:cs" if "name:cs" in df.columns else "name"
        if name_col not in df.columns:
            raise ValueError("V CSV chybí sloupec 'name:cs' i 'name'.")
        mask = df[name_col].astype(str).str.strip() == target_product_name.strip()
        if mask.any():
            return mask
        raise ValueError(f"Produkt '{target_product_name}' nebyl nalezen.")

    return df.index == 0


def ensure_output_columns(df: pd.DataFrame) -> pd.DataFrame:
    required_cols = [
        "name:cs", "name:en", "name:sk",
        "shortDescription:cs", "shortDescription:en", "shortDescription:sk",
        "description:cs", "description:en", "description:sk",
        "xmlFeedName:cs", "xmlFeedName:en", "xmlFeedName:sk",
        "seoTitle:cs", "seoTitle:en", "seoTitle:sk",
        "metaDescription:cs", "metaDescription:en", "metaDescription:sk",
    ]
    for col in required_cols:
        if col not in df.columns:
            df[col] = ""
    return df


def build_create_csv(source_row: pd.Series, final_name_cs: str, final_detail_html_cs: str) -> pd.DataFrame:
    create_columns = [
        "code",
        "pairCode",
        "name",
        "price",
        "description",
        "image", "image2", "image3", "image4", "image5",
        "image6", "image7", "image8", "image9", "image10",
        "image11", "image12", "image13", "image14", "image15",
        "image16", "image17", "image18", "image19", "image20",
    ]

    create_data = {
        "code": [str(source_row.get("code", "")).strip()],
        "pairCode": [str(source_row.get("pairCode", "")).strip()],
        "name": [str(final_name_cs).strip()],
        "price": [str(source_row.get("price", "")).strip()],
        "description": [str(final_detail_html_cs).strip()],
    }

    for col in create_columns[5:]:
        create_data[col] = [str(source_row.get(col, "")).strip()]

    return pd.DataFrame(create_data, columns=create_columns)


# =========================
# RUN
# =========================

def run_filler(
    template_type: str,
    csv_path: str,
    template_dir: str,
    prompt_output_docx_path: str,
    output_csv_path: str,
    output_create_csv_path: str,
    target_product_name: Optional[str] = None,
    target_ean: Optional[str] = None,
    debug: bool = True,
    extra_values: Optional[Dict[str, str]] = None,
) -> Dict[str, object]:
    template_type = template_type.strip().lower()
    csv_path = Path(csv_path).expanduser()
    template_dir = Path(template_dir).expanduser()
    prompt_output_docx_path = Path(prompt_output_docx_path).expanduser()
    output_csv_path = Path(output_csv_path).expanduser()
    output_create_csv_path = Path(output_create_csv_path).expanduser()

    template_paths = get_template_paths(template_dir, template_type)
    df = pd.read_csv(csv_path, sep=";", dtype=str).fillna("")
    df = ensure_output_columns(df)

    mask = build_csv_row_mask(df, target_product_name=target_product_name, target_ean=target_ean)
    matched_rows = df.loc[mask]

    if matched_rows.empty:
        raise ValueError("Nepodařilo se najít cílový řádek v CSV.")

    source_row = matched_rows.iloc[0]

    name_col = "name:cs" if "name:cs" in df.columns and str(source_row.get("name:cs", "")).strip() else "name"
    product_name = str(source_row.get(name_col, "")).strip()

    short_template_cs = read_docx_text(template_paths["short_cs"])
    short_template_en = read_docx_text(template_paths["short_en"])
    short_template_sk = read_docx_text(template_paths["short_sk"])

    detail_template_cs = read_docx_text(template_paths["detail_cs"])
    detail_template_en = read_docx_text(template_paths["detail_en"])
    detail_template_sk = read_docx_text(template_paths["detail_sk"])

    short_placeholders_cs = extract_placeholders(short_template_cs)
    detail_placeholders_cs = extract_placeholders(detail_template_cs)

    prompt_output_text = read_docx_text(prompt_output_docx_path)
    values_by_lang = parse_prompt_output_by_lang(prompt_output_text)

    if extra_values:
        normalized_extra = {
            canonical_key(k): normalize_docx_text(v).strip()
            for k, v in extra_values.items()
        }
        for lang in ("cs", "en", "sk"):
            values_by_lang[lang].update(normalized_extra)

    for lang in ("cs", "en", "sk"):
        if "sekce_o_jednotce_nadpis" not in values_by_lang[lang] and "section_inside_title" in values_by_lang[lang]:
            values_by_lang[lang]["sekce_o_jednotce_nadpis"] = values_by_lang[lang]["section_inside_title"]

        if "sekce_o_jednotce_text" not in values_by_lang[lang] and "section_inside_text" in values_by_lang[lang]:
            values_by_lang[lang]["sekce_o_jednotce_text"] = values_by_lang[lang]["section_inside_text"]

    for lang in ("cs", "en", "sk"):
        current_product_name = product_name
        if values_by_lang[lang].get(canonical_key("nazev_produktu"), "").strip():
            current_product_name = values_by_lang[lang][canonical_key("nazev_produktu")]

        values_by_lang[lang] = inject_safe_defaults(values_by_lang[lang], current_product_name)
        values_by_lang[lang] = apply_aliases(values_by_lang[lang], template_type)

    if debug:
        print("PROMPT FILE:", prompt_output_docx_path)
        print("DETAIL FILE:", template_paths["detail_cs"])
        print()
        print_placeholder_summary("KRÁTKÁ ŠABLONA CS", short_placeholders_cs, values_by_lang["cs"])
        print_placeholder_summary("DETAILNÍ ŠABLONA CS", detail_placeholders_cs, values_by_lang["cs"])

    final_short_html_cs = flatten_html_for_csv(
        fill_template(short_template_cs, values_by_lang["cs"], strict=True)
    )
    final_short_html_en = flatten_html_for_csv(
        fill_template(short_template_en, values_by_lang["en"], strict=True)
    )
    final_short_html_sk = flatten_html_for_csv(
        fill_template(short_template_sk, values_by_lang["sk"], strict=True)
    )

    final_detail_html_cs = flatten_html_for_csv(
        fill_template(detail_template_cs, values_by_lang["cs"], strict=True)
    )
    final_detail_html_en = flatten_html_for_csv(
        fill_template(detail_template_en, values_by_lang["en"], strict=True)
    )
    final_detail_html_sk = flatten_html_for_csv(
        fill_template(detail_template_sk, values_by_lang["sk"], strict=True)
    )

    df.loc[mask, "shortDescription:cs"] = final_short_html_cs
    df.loc[mask, "shortDescription:en"] = final_short_html_en
    df.loc[mask, "shortDescription:sk"] = final_short_html_sk

    df.loc[mask, "description:cs"] = final_detail_html_cs
    df.loc[mask, "description:en"] = final_detail_html_en
    df.loc[mask, "description:sk"] = final_detail_html_sk

    source_name_base = str(source_row.get("name", "")).strip()

    final_name_cs = str(source_row.get("name:cs", "")).strip() or source_name_base
    final_name_en = str(source_row.get("name:en", "")).strip() or source_name_base
    final_name_sk = str(source_row.get("name:sk", "")).strip() or source_name_base

    df.loc[mask, "name:cs"] = final_name_cs
    df.loc[mask, "name:en"] = final_name_en
    df.loc[mask, "name:sk"] = final_name_sk

    final_xml_feed_name_cs = final_name_cs
    final_xml_feed_name_en = final_name_en
    final_xml_feed_name_sk = final_name_sk

    df.loc[mask, "xmlFeedName:cs"] = final_xml_feed_name_cs
    df.loc[mask, "xmlFeedName:en"] = final_xml_feed_name_en
    df.loc[mask, "xmlFeedName:sk"] = final_xml_feed_name_sk

    df.loc[mask, "seoTitle:cs"] = f"{final_name_cs} | PlasticWargaming" if final_name_cs else ""
    df.loc[mask, "seoTitle:en"] = f"{final_name_en} | PlasticWargaming" if final_name_en else ""
    df.loc[mask, "seoTitle:sk"] = f"{final_name_sk} | PlasticWargaming" if final_name_sk else ""

    meta_key = canonical_key("doporučení_pro_koho")

    dop_cs = values_by_lang["cs"].get("doporučení_pro_koho", "").strip() or values_by_lang["cs"].get(meta_key, "").strip()
    dop_en = values_by_lang["en"].get("doporučení_pro_koho", "").strip() or values_by_lang["en"].get(meta_key, "").strip()
    dop_sk = values_by_lang["sk"].get("doporučení_pro_koho", "").strip() or values_by_lang["sk"].get(meta_key, "").strip()

    df.loc[mask, "metaDescription:cs"] = dop_cs if dop_cs else ""
    df.loc[mask, "metaDescription:en"] = dop_en if dop_en else ""
    df.loc[mask, "metaDescription:sk"] = dop_sk if dop_sk else ""

    for col in ["name", "xmlFeedName", "seoTitle", "system", "faction", "productType", "hp_url", "gw_url"]:
        if col in df.columns:
            df = df.drop(columns=[col])

    output_csv_path.parent.mkdir(parents=True, exist_ok=True)
    output_create_csv_path.parent.mkdir(parents=True, exist_ok=True)

    df.to_csv(
        output_csv_path,
        sep=";",
        index=False,
        encoding="utf-8-sig",
        quoting=csv.QUOTE_ALL
    )

    create_source_row = source_row.copy()
    create_source_row["image"] = extra_values.get("intro_image_src", source_row.get("image", "")) if extra_values else source_row.get("image", "")
    create_source_row["image2"] = extra_values.get("img1_src", source_row.get("image2", "")) if extra_values else source_row.get("image2", "")
    create_source_row["image3"] = extra_values.get("img2_src", source_row.get("image3", "")) if extra_values else source_row.get("image3", "")
    create_source_row["image4"] = extra_values.get("img3_src", source_row.get("image4", "")) if extra_values else source_row.get("image4", "")
    create_source_row["image5"] = extra_values.get("img4_src", source_row.get("image5", "")) if extra_values else source_row.get("image5", "")

    create_df = build_create_csv(
        source_row=create_source_row,
        final_name_cs=final_name_cs,
        final_detail_html_cs=final_detail_html_cs,
    )

    create_df.to_csv(
        output_create_csv_path,
        sep=";",
        index=False,
        encoding="utf-8-sig",
        quoting=csv.QUOTE_ALL
    )

    print("=" * 70)
    print("HOTOVO")
    print("=" * 70)
    print(f"TEMPLATE_TYPE: {template_type}")
    print(f"Produkt: {product_name}")
    print(f"MULTI CSV uloženo do:\n{output_csv_path}")
    print(f"CREATE CSV uloženo do:\n{output_create_csv_path}")

    return {
        "template_type": template_type,
        "product_name": product_name,
        "output_csv": str(output_csv_path),
        "output_create_csv": str(output_create_csv_path),
        "final_name_cs": final_name_cs,
        "short_html_cs": final_short_html_cs,
        "detail_html_cs": final_detail_html_cs,
    }


# =========================
# MAIN
# =========================

def main() -> None:
    template_type = "miniatures"
    csv_path = "/Users/janholubik/Downloads/example_SOURCE.csv"
    template_dir = "/Users/janholubik/Downloads/XML/XML_plastic/sablony"
    prompt_output_docx_path = "/Users/janholubik/Downloads/vystup_prompt.docx"
    output_csv_path = "/Users/janholubik/Downloads/0_FILLED.csv"
    output_create_csv_path = "/Users/janholubik/Downloads/0_CREATE.csv"

    run_filler(
        template_type=template_type,
        csv_path=csv_path,
        template_dir=template_dir,
        prompt_output_docx_path=prompt_output_docx_path,
        output_csv_path=output_csv_path,
        output_create_csv_path=output_create_csv_path,
        target_product_name=None,
        target_ean=None,
        debug=True,
        extra_values=None,
    )


if __name__ == "__main__":
    main()