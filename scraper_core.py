#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import csv
import re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from urllib.parse import urljoin, urlparse, urlunparse

import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document


# =========================
# OUTPUT COLUMNS
# =========================
CREATE_COLUMNS = [
    "code",
    "pairCode",
    "name",
    "price",
    "description",
    "image", "image2", "image3", "image4", "image5",
    "image6", "image7", "image8", "image9", "image10",
    "image11", "image12", "image13", "image14", "image15",
    "image16", "image17", "image18", "image19", "image20",
]

SOURCE_COLUMNS = [
    "code",
    "pairCode",
    "name",
    "ean",
    "externalCode",
    "price",
    "priceWithoutVat",
    "standardPrice",
    "categoryText",
    "warranty",
    "supplier",
    "googleCategoryIdInFeed",
    "heurekaCategoryId",
    "zboziCategoryId",
    "googleCategoryId",
    "stock",
    "percentVat",
    "ossTaxRate:CZ",
    "availabilityInStock",
    "availabilityOutOfStock",
    "productVisibility",
    "xmlFeedName",
    "seoTitle",
    "system",
    "faction",
    "productType",
    "hp_url",
    "gw_url",
    "image", "image2", "image3", "image4", "image5",
    "image6", "image7", "image8", "image9", "image10",
    "image11", "image12", "image13", "image14", "image15",
    "image16", "image17", "image18", "image19", "image20",
]

BASE_HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X) AppleWebKit/537.36 (KHTML, like Gecko) Chrome Safari",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
    "Accept-Language": "cs-CZ,cs;q=0.9,en;q=0.8",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
}

GW_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/122.0.0.0 Safari/537.36"
    ),
    "Accept": (
        "text/html,application/xhtml+xml,application/xml;q=0.9,"
        "image/avif,image/webp,image/apng,*/*;q=0.8,"
        "application/signed-exchange;v=b3;q=0.7"
    ),
    "Accept-Language": "en-US,en;q=0.9",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Upgrade-Insecure-Requests": "1",
    "Sec-Fetch-Dest": "document",
    "Sec-Fetch-Mode": "navigate",
    "Sec-Fetch-Site": "same-origin",
    "Sec-Fetch-User": "?1",
}

FX_RATES = {
    "€": 24.50,
    "$": 21.00,
    "£": 28.00,
}

FACTIONS_BY_SYSTEM = {
    "Warhammer 40k": [
        "Black Templars", "Blood Angels", "Dark Angels", "Deathwatch", "Grey Knights",
        "Imperial Fists", "Iron Hands", "Raven Guard", "Salamanders", "Space Wolves",
        "Ultramarines", "White Scars", "Space Marines", "Astra Militarum",
        "Adepta Sororitas", "Adeptus Mechanicus", "Imperial Knights", "Adeptus Custodes",
        "Imperial Agents", "Death Guard", "Thousand Sons", "World Eaters",
        "Chaos Space Marines", "Chaos Daemons", "Chaos Knights", "Traitor Guard",
        "Emperor's Children", "Aeldari", "Drukhari", "Orks", "Tyranids",
        "Genestealer Cults", "T'au Empire", "Necrons", "Leagues of Votann",
    ],
    "Warhammer Age of Sigmar": [
        "Stormcast Eternals", "Fyreslayers", "Kharadron Overlords", "Sylvaneth",
        "Lumineth Realm-lords", "Cities of Sigmar", "Daughters of Khaine",
        "Idoneth Deepkin", "Seraphon", "Slaves to Darkness", "Blades of Khorne",
        "Disciples of Tzeentch", "Maggotkin of Nurgle", "Hedonites of Slaanesh",
        "Skaven", "Soulblight Gravelords", "Ossiarch Bonereapers", "Nighthaunt",
        "Flesh-eater Courts", "Orruk Warclans", "Ogor Mawtribes", "Gloomspite Gitz",
        "Sons of Behemat",
    ],
    "The Horus Heresy": [
        "Dark Angels", "White Scars", "Space Wolves", "Imperial Fists", "Blood Angels",
        "Iron Hands", "Ultramarines", "Salamanders", "Raven Guard", "Knights-Errant",
        "Sons of Horus", "Emperor's Children", "Death Guard", "World Eaters",
        "Thousand Sons", "Word Bearers", "Night Lords", "Iron Warriors",
        "Alpha Legion", "Cult Mechanicum", "Talons of the Emperor", "Solar Auxilia",
    ],
    "Warhammer: The Old World": [
        "Grand Cathay", "Wood Elf Realms", "Kingdom of Bretonnia", "High Elf Realms",
        "Empire of Man", "Dwarfen Mountain Holds", "Warriors of Chaos",
        "Tomb Kings of Khemri", "Orcs and Goblin Tribes", "Beastmen Brayherd",
    ],
}

FACTION_ALIASES = {
    "Aeldari": ["aeldari", "eldar"],
    "Drukhari": ["drukhari", "dark eldar"],
    "Adepta Sororitas": ["adepta sororitas", "sisters of battle"],
    "Astra Militarum": ["astra militarum", "imperial guard"],
    "Imperial Knights": ["imperial knights", "questor imperialis"],
    "Space Marines": ["space marines", "adeptus astartes"],
    "Emperor's Children": ["emperor's children", "emperors children"],
    "T'au Empire": ["t'au empire", "tau empire"],
    "Ogor Mawtribes": ["ogor mawtribes", "ogres"],
    "Flesh-eater Courts": ["flesh-eater courts", "flesh eater courts"],
    "Lumineth Realm-lords": ["lumineth realm-lords", "lumineth realm lords"],
}


def log(msg: str, verbose: bool) -> None:
    if verbose:
        print(msg)


# =========================
# LINKS LOADER
# =========================
def load_links_raw(path: str) -> pd.DataFrame:
    p = Path(path).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"Soubor neexistuje: {p}")

    if p.suffix.lower() in (".txt", ".list"):
        rows = []
        with p.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#"):
                    continue
                low = line.lower().replace(" ", "")
                if low.startswith("hp_url") and ("gw_url" in low):
                    continue

                parts = [x.strip() for x in re.split(r"[;\t,]+", line) if x.strip()]
                if not parts:
                    continue

                hp = parts[0]
                gw = parts[1] if len(parts) > 1 else ""
                if not hp.startswith("http"):
                    continue
                if gw and not gw.startswith("http"):
                    gw = ""

                rows.append({"hp_url": hp, "gw_url": gw})

        if not rows:
            raise ValueError("TXT soubor je prázdný nebo neobsahuje validní URL řádky.")
        return pd.DataFrame(rows, columns=["hp_url", "gw_url"]).fillna("")

    last_err = None
    for sep in (";", ","):
        try:
            df = pd.read_csv(p, sep=sep, encoding="utf-8-sig", dtype=str).fillna("")
            df.columns = [str(c).strip() for c in df.columns]
            return df
        except Exception as e:
            last_err = e

    raise ValueError(f"Nepodařilo se načíst soubor jako CSV/TXT. Poslední chyba: {last_err}")


def sanitize_links_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    if "hp_url" not in df.columns and len(df.columns) >= 1:
        df = df.rename(columns={df.columns[0]: "hp_url"})
    if "gw_url" not in df.columns:
        if len(df.columns) >= 2:
            df = df.rename(columns={df.columns[1]: "gw_url"})
        else:
            df["gw_url"] = ""

    df["hp_url"] = df["hp_url"].astype(str).str.strip()
    df["gw_url"] = df["gw_url"].astype(str).str.strip()

    df = df[df["hp_url"].str.lower().ne("hp_url")]
    df = df[df["hp_url"].str.startswith("http")]
    df.loc[~df["gw_url"].str.startswith("http"), "gw_url"] = ""

    df = df.reset_index(drop=True)
    if df.empty:
        raise ValueError("Po vyčištění nezbyl žádný validní řádek s hp_url.")

    return df[["hp_url", "gw_url"]]


def load_links(path: str) -> pd.DataFrame:
    return sanitize_links_df(load_links_raw(path))


# =========================
# HELPERS
# =========================
def fetch_html(
    url: str,
    timeout: int = 30,
    referer: str = "",
    headers: Optional[Dict[str, str]] = None
) -> Tuple[str, str]:
    hdrs = dict(headers or BASE_HEADERS)

    if referer:
        hdrs["Referer"] = referer

    session = requests.Session()

    r = session.get(
        url,
        headers=hdrs,
        timeout=timeout,
        allow_redirects=True,
    )
    r.raise_for_status()
    return (r.text or ""), str(r.url)


def strip_query_and_fragment(url: str) -> str:
    p = urlparse(url)
    return urlunparse((p.scheme, p.netloc, p.path, "", "", ""))


def looks_like_gw_product_html(html: str) -> bool:
    if not html:
        return False
    markers = [
        'data-testid="hero-product-card-price"',
        'data-testid="quantity-and-price-container"',
        'data-testid="gallery-modal-image"',
        'data-testid="button-gallery-view-full"',
        'data-testid="image-carousel-image-button"',
        'data-testid="gallery-image-button"',
        'data-testid="image-gallery"',
        "/app/resources/catalog/product/",
        "Missing_Image_Servo_Skull",
    ]
    low = html.lower()
    return any(m.lower() in low for m in markers)


def fetch_gw_html(url: str, timeout: int = 30, verbose: bool = False) -> Tuple[str, str]:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Cache-Control": "no-cache",
        "Pragma": "no-cache",
        "Upgrade-Insecure-Requests": "1",
        "Referer": "https://www.google.com/",
    }

    session = requests.Session()
    resp = session.get(
        url,
        headers=headers,
        timeout=timeout,
        allow_redirects=True,
    )

    html = resp.text or ""
    final_url = str(resp.url)

    if verbose:
        print(f"GW status: {resp.status_code}")
        print(f"GW final URL: {final_url}")
        print(f"GW HTML length: {len(html)}")
        print('has hero price:', 'data-testid="hero-product-card-price"' in html)
        print('has quantity price:', 'data-testid="quantity-and-price-container"' in html)
        print('has image gallery:', 'data-testid="image-gallery"' in html)
        print('has gallery button:', 'data-testid="gallery-image-button"' in html)
        print('has carousel button:', 'data-testid="image-carousel-image-button"' in html)
        print('has product path:', '/app/resources/catalog/product/' in html)

    with open("DEBUG_GW.html", "w", encoding="utf-8") as f:
        f.write(html)

    return html, final_url


def norm_ws(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip()


def clean_title(title: str) -> str:
    title = norm_ws(title)
    title = re.sub(r"^Knihy\s*-\s*", "", title, flags=re.IGNORECASE)
    return title


def to_float_price(txt: str) -> Optional[float]:
    if not txt:
        return None
    t = txt.replace("\xa0", " ")
    t = re.sub(r"[^\d,\. ]", "", t)
    t = t.replace(" ", "")

    if t.count(",") == 1 and t.count(".") >= 1:
        t = t.replace(".", "")
        t = t.replace(",", ".")
    else:
        if t.count(",") == 1 and t.count(".") == 0:
            t = t.replace(",", ".")
        if t.count(".") == 1 and t.count(",") == 0:
            parts = t.split(".")
            if len(parts[-1]) == 3:
                t = "".join(parts)

    try:
        return float(t)
    except Exception:
        return None


def price_without_vat(price_with_vat: float, vat_percent: float) -> float:
    return round(price_with_vat / (1.0 + vat_percent / 100.0), 2)


def load_docx_raw_text(docx_path: str) -> str:
    doc = Document(docx_path)
    parts: List[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join([x for x in parts if x is not None])


def safe_singleline_html(s: str) -> str:
    s = (s or "").replace("\r\n", "\n").replace("\r", "\n")
    return s.replace("\n", "<br>")


def normalize_system_name(system: str) -> str:
    system = norm_ws(system)
    if system == "Warhammer 40,000":
        return "Warhammer 40k"
    return system


def slugify_filename(s: str, max_len: int = 120) -> str:
    s = norm_ws(s)
    if not s:
        return "produkt"
    s = s.replace("/", " ").replace("\\", " ").replace(":", " ").replace("|", " ")
    s = re.sub(r"[^\w\-. ]+", "", s, flags=re.UNICODE)
    s = re.sub(r"\s+", "_", s).strip("_")
    if not s:
        s = "produkt"
    return s[:max_len]


def make_xml_feed_name_from_name(name: str) -> str:
    value = clean_title(norm_ws(name))
    return value if value else "Produkt"


def make_seo_title_from_name(name: str) -> str:
    value = make_xml_feed_name_from_name(name)
    return f"{value} | PlasticWargaming"


def fmt_cz_money(v: Optional[float], decimals: int = 2) -> str:
    if v is None:
        return "-"
    return f"{v:.{decimals}f}".replace(".", ",")

def is_valid_product_code(value: str) -> bool:
    v = norm_ws(value)

    if not v:
        return False

    # moc krátké nesmysly typu "ze"
    if len(v) < 5:
        return False

    # čistě písmena bez čísel/hyfenů bývají často špatně
    if re.fullmatch(r"[a-zA-Z]+", v):
        return False

    return True


def generate_fallback_code_from_url(hp_url: str, gw_url: str = "") -> str:
    source_url = gw_url or hp_url
    path = urlparse(source_url).path.strip("/")

    slug = path.split("/")[-1] if path else "produkt"
    slug = re.sub(r"[^a-zA-Z0-9\-]", "-", slug)
    slug = re.sub(r"-+", "-", slug).strip("-").lower()

    if not slug:
        slug = "produkt"

    return f"wh-{slug}"


# =========================
# TYPE DETECTION
# =========================
def detect_product_type(title_raw: str, category_text: str, ean: str) -> str:
    s = f"{title_raw} {category_text}".lower()
    e = str(ean or "")

    if e.startswith("978") or any(k in s for k in [
        "codex", "rulebook", "battletome", "hardback", "paperback",
        "black library", "novel", "book"
    ]):
        return "book"

    if any(k in s for k in [
        "dice", "dice set", "kostk", "kub", "d6", "d10", "d20"
    ]):
        return "dice"

    if any(k in s for k in [
        "warscroll", "warscoll",
        "datacard", "data card",
        "reference cards", "reference card pack",
        "card pack", "cards", "karty"
    ]):
        return "warscroll"

    if any(k in s for k in [
        "upgrade", "upgrades", "upgrade pack", "upgrade set", "weapon upgrade",
        "head upgrade", "shoulder pad upgrade", "upgrade kit", "upgrade sprue",
        "upgrades and transfers", "upgrade & transfers", "transfer sheet",
    ]):
        return "upgrades"

    if any(k in s for k in [
        "terrain", "scenery", "barricades", "ruins", "paint", "paints", "brush",
        "brushes", "glue", "primer", "spray", "tool", "tools", "basing", "tuft",
        "texture", "palette", "water pot", "mouldline remover", "hobby knife",
        "clippers", "files", "drill", "paint set", "gaming aids", "token set"
    ]):
        return "accessories"

    return "miniatures"


# =========================
# HP SCRAPE
# =========================
def hp_extract_h1(soup: BeautifulSoup) -> str:
    h1 = soup.select_one("h1")
    return norm_ws(h1.get_text(" ")) if h1 else ""


def hp_extract_price(soup: BeautifulSoup) -> Optional[float]:
    selectors = [
        "strong.price-final[data-testid='productCardPrice'] span.price-final-holder",
        "span.price-final-holder",
        "strong.price-final",
        "[data-testid='productPrice']",
        ".price-final",
        ".product-price",
    ]
    for sel in selectors:
        el = soup.select_one(sel)
        if el:
            p = to_float_price(el.get_text(" "))
            if p is not None:
                return p
    txt = soup.get_text("\n")
    m = re.search(r"(\d[\d\s\.\,]+)\s*Kč", txt)
    return to_float_price(m.group(0)) if m else None


def hp_extract_breadcrumbs_list(soup: BeautifulSoup) -> List[str]:
    root = soup.select_one('div.breadcrumbs[itemscope][itemtype="https://schema.org/BreadcrumbList"]')
    if root:
        spans = root.select('[itemprop="itemListElement"] [itemprop="name"]')
    else:
        spans = soup.select('[itemtype="https://schema.org/BreadcrumbList"] [itemprop="itemListElement"] [itemprop="name"]')

    items: List[str] = []
    for sp in spans:
        t = norm_ws(sp.get_text(" "))
        if not t:
            continue
        low = t.lower()
        if low in ("domů", "home", "herní prostor", "games workshop"):
            continue
        items.append(t)

    out, seen = [], set()
    for x in items:
        if x in seen:
            continue
        seen.add(x)
        out.append(x)
    return out


def hp_pick_system(soup: BeautifulSoup) -> str:
    crumbs = hp_extract_breadcrumbs_list(soup)

    for c in crumbs:
        if c in ("Warhammer 40,000", "The Horus Heresy"):
            return normalize_system_name(c)
        if c.startswith("Warhammer Age of Sigmar"):
            return normalize_system_name(c)

    for c in crumbs:
        if c.lower() != "warhammer":
            return normalize_system_name(c)

    return normalize_system_name("Warhammer")


def hp_extract_ean(soup: BeautifulSoup) -> str:
    txt = soup.get_text("\n")
    m = re.search(r"\bEAN\b[^\d]*(\d{8,14})\b", txt, flags=re.IGNORECASE)
    if m:
        return m.group(1)

    m2 = re.search(r"\b(\d{13})\b", txt)
    if m2:
        return m2.group(1)

    for script in soup.select("script[type='application/ld+json']"):
        txt_json = script.get_text(" ", strip=True)
        m3 = re.search(r'"gtin13"\s*:\s*"(\d{13})"', txt_json, flags=re.I)
        if m3:
            return m3.group(1)
        m4 = re.search(r'"gtin"\s*:\s*"(\d{8,14})"', txt_json, flags=re.I)
        if m4:
            return m4.group(1)

    return ""


def hp_extract_external_code(soup: BeautifulSoup) -> str:
    txt = soup.get_text("\n")

    m = re.search(r"\b(\d{2,3}\-\d{2,4})\b", txt)
    if m:
        return m.group(1)

    m2 = re.search(
        r"(Kód|Code|Produktové číslo|Číslo produktu|SKU)[^A-Za-z0-9]{0,10}([A-Za-z0-9][A-Za-z0-9\-\/]{4,})",
        txt,
        flags=re.IGNORECASE
    )
    if m2:
        candidate = m2.group(2).strip()
        if is_valid_product_code(candidate):
            return candidate

    for script in soup.select("script[type='application/ld+json']"):
        txt_json = script.get_text(" ", strip=True)
        m3 = re.search(r'"sku"\s*:\s*"([^"]+)"', txt_json, flags=re.I)
        if m3:
            return m3.group(1).strip()

    return ""


def build_name_from_h1(h1: str, system: str = "", faction: str = "") -> str:
    h1 = clean_title(h1)
    system = normalize_system_name(system)
    faction = norm_ws(faction)

    if not h1:
        return "Produkt"

    if system and faction:
        prefix = f"{system}: {faction}"
        if not h1.lower().startswith(prefix.lower()):
            return f"{prefix} - {h1}"
        return h1

    if system:
        if not h1.lower().startswith(system.lower()):
            return f"{system}: {h1}"
        return h1

    return h1


def extract_relevant_gw_faction_text(gw_html: str) -> str:
    if not gw_html:
        return ""

    soup = BeautifulSoup(gw_html, "html.parser")
    parts: List[str] = []

    for el in soup.select('[data-testid^="product-detail-feature-"]'):
        txt = norm_ws(el.get_text(" "))
        if txt:
            parts.append(txt)

    for el in soup.select('[data-testid="hero-product-card-name"]'):
        txt = norm_ws(el.get_text(" "))
        if txt:
            parts.append(txt)

    if soup.title:
        txt = norm_ws(soup.title.get_text(" "))
        if txt:
            parts.append(txt)

    for el in soup.select('link[rel="canonical"], meta[property="og:url"]'):
        u = (el.get("href") or el.get("content") or "").strip()
        if u:
            parts.append(u)

    return " | ".join(parts)


def detect_faction_from_gw_html(gw_html: str, system: str = "") -> str:
    if not gw_html:
        return ""

    text = extract_relevant_gw_faction_text(gw_html)
    text_low = text.lower()

    faction_list = FACTIONS_BY_SYSTEM.get(system, [])
    if not faction_list:
        return ""

    matches: List[str] = []

    for faction in faction_list:
        aliases = FACTION_ALIASES.get(faction, [faction.lower()])
        for alias in aliases:
            alias_low = alias.lower()
            if alias_low in text_low:
                matches.append(faction)
                break

    if not matches:
        return ""

    matches = sorted(set(matches), key=len, reverse=True)
    return matches[0]


# =========================
# GW PRICE
# =========================
def gw_extract_price(gw_html: str) -> Tuple[Optional[float], Optional[str]]:
    if not gw_html:
        return None, None

    try:
        soup = BeautifulSoup(gw_html, "lxml")

        hero = soup.select_one('[data-testid="hero-product-card-price"]')
        if hero:
            txt = hero.get_text(" ", strip=True)
            m = re.search(r"([€$£])\s*([0-9]+(?:[.,][0-9]{1,2})?)", txt)
            if m:
                return float(m.group(2).replace(",", ".")), m.group(1)

        qty = soup.select_one('[data-testid="quantity-and-price-container"]')
        if qty:
            txt = qty.get_text(" ", strip=True)
            m = re.search(r"([€$£])\s*([0-9]+(?:[.,][0-9]{1,2})?)", txt)
            if m:
                return float(m.group(2).replace(",", ".")), m.group(1)
    except Exception:
        pass

    patterns = [
        r'data-testid="hero-product-card-price".{0,600}?([€$£])\s*([0-9]+(?:[.,][0-9]{1,2})?)',
        r'data-testid="quantity-and-price-container".{0,600}?([€$£])\s*([0-9]+(?:[.,][0-9]{1,2})?)',
        r'([€$£])\s*([0-9]+(?:[.,][0-9]{1,2})?)',
    ]
    for pat in patterns:
        m = re.search(pat, gw_html, flags=re.I | re.S)
        if m:
            try:
                return float(m.group(2).replace(",", ".")), m.group(1)
            except Exception:
                pass

    return None, None


# =========================
# GW IMAGE SCRAPE
# =========================
def abs_url(base: str, u: str) -> str:
    return urljoin(base, u)


def strip_query(u: str) -> str:
    try:
        p = urlparse(u)
        return urlunparse((p.scheme, p.netloc, p.path, "", "", ""))
    except Exception:
        return u


def filename_key(u: str) -> str:
    u2 = strip_query(u)
    return Path(urlparse(u2).path).name.lower()


def pick_best_srcset(srcset: str) -> Optional[str]:
    if not srcset:
        return None
    parts = []
    for p in srcset.split(","):
        p = p.strip()
        if not p:
            continue
        parts.append(p.split(" ")[0].strip())
    return parts[-1] if parts else None


def extract_imgs_from_node(node: BeautifulSoup, base_url: str) -> List[str]:
    urls: List[str] = []

    if node is None:
        return urls

    node_name = getattr(node, "name", None)

    # Když je přímo <img>
    if node_name == "img":
        src = (node.get("src") or "").strip()
        if src:
            urls.append(abs_url(base_url, src))

        best = pick_best_srcset((node.get("srcset") or "").strip())
        if best:
            urls.append(abs_url(base_url, best))

    # Když je přímo <source>
    if node_name == "source":
        best = pick_best_srcset((node.get("srcset") or "").strip())
        if best:
            urls.append(abs_url(base_url, best))

    # všechny vnořené img
    for img in node.select("img"):
        src = (img.get("src") or "").strip()
        if src:
            urls.append(abs_url(base_url, src))

        best = pick_best_srcset((img.get("srcset") or "").strip())
        if best:
            urls.append(abs_url(base_url, best))

    # všechny vnořené source
    for s in node.select("source"):
        best = pick_best_srcset((s.get("srcset") or "").strip())
        if best:
            urls.append(abs_url(base_url, best))

    return urls


def is_missing_image(u: str) -> bool:
    return "Missing_Image_Servo_Skull" in u or "missing_image" in u.lower()


def filter_gw_product_images(urls: List[str], keep_360: bool) -> List[str]:
    out: List[str] = []

    for u in urls:
        if not u or is_missing_image(u):
            continue

        try:
            p = urlparse(u)
            path = (p.path or "").lower()
            netloc = (p.netloc or "").lower()
        except Exception:
            path = str(u).lower()
            netloc = ""

        # POZOR:
        # nepoužívat jen "360" in path
        # protože GW má 360 i uvnitř názvů souborů (např. AoSWarcry18360)
        is_real_360_path = (
            "/threesixty/" in path
            or "/threeSixty/" in (p.path or "")
        )

        if (not keep_360) and is_real_360_path:
            continue

        allowed = (
            "warhammer.com" in netloc
            or "games-workshop.com" in netloc
            or "gwplc" in netloc
            or "/app/resources/catalog/product/" in path
            or "/catalog/product/" in path
            or "/media/catalog/product/" in path
            or "/resources/catalog/product/" in path
        )

        looks_like_image = any(ext in path for ext in [".jpg", ".jpeg", ".png", ".webp", ".avif"])

        if allowed or looks_like_image:
            out.append(u)

    return out


def uniq_keep_order(seq: List[str]) -> List[str]:
    seen = set()
    out = []
    for x in seq:
        if x and x not in seen:
            out.append(x)
            seen.add(x)
    return out


def dedupe_by_filename(urls: List[str]) -> List[str]:
    seen = set()
    out = []
    for u in urls:
        key = filename_key(u)
        if key and key not in seen:
            out.append(u)
            seen.add(key)
    return out


def expected_count_from_html(html: str) -> Optional[int]:
    m = re.findall(r"image\s+\d+\s+of\s+(\d+)", html, flags=re.I)
    nums = []
    for x in m:
        try:
            nums.append(int(x))
        except Exception:
            pass
    return max(nums) if nums else None


def has_full_gallery_button(soup: BeautifulSoup) -> bool:
    return (
        soup.select_one("#gallery-view-full") is not None
        or soup.select_one('[data-testid="button-gallery-view-full"]') is not None
        or soup.select_one('[data-testid="button-gallery-view-full"] p') is not None
    )


def ensure_query_defaults(url: str, default_q: str) -> str:
    if not url:
        return url
    try:
        p = urlparse(url)
        if p.query:
            return url
        if default_q.startswith("?"):
            default_q = default_q[1:]
        return urlunparse((p.scheme, p.netloc, p.path, p.params, default_q, p.fragment))
    except Exception:
        return url
    
def scrape_gw_images_stable(
    gw_url: str,
    html: str,
    max_images: int = 20,
    keep_360: bool = False,
    ensure_query: bool = False,
    ensure_query_default: str = "fm=webp&w=1200&h=1237",
) -> List[str]:
    soup = BeautifulSoup(html, "lxml")
    urls: List[str] = []

    # 1) Nejdřív zkus přímo full gallery
    for gallery in soup.select('[data-testid="image-gallery"]'):
        urls.extend(extract_imgs_from_node(gallery, gw_url))

    # 2) Pak jednotlivé gallery itemy
    for node in soup.select(
        '[data-testid="gallery-image"], '
        '[data-testid="gallery-image-button"]'
    ):
        urls.extend(extract_imgs_from_node(node, gw_url))

    # 3) Pak carousel jako fallback
    for node in soup.select(
        '[data-testid="container-image-carousel"], '
        '[data-testid="image-carousel"], '
        '[data-testid="image-carousel-image-button"], '
        '[data-testid="image-carousel-active-image"], '
        '[data-testid="image-carousel-mobile"], '
        '[data-testid="image-carousel-mobile-image"]'
    ):
        urls.extend(extract_imgs_from_node(node, gw_url))

    # 4) Poslední nouzový fallback — všechny img/source z dokumentu,
    # které vypadají jako produktové obrázky
    if not urls:
        for el in soup.select(
            'source[srcset*="/catalog/product/"], '
            'source[srcset*="/app/resources/catalog/product/"], '
            'img[src*="/catalog/product/"], '
            'img[src*="/app/resources/catalog/product/"]'
        ):
            if el.name == "source":
                best = pick_best_srcset((el.get("srcset") or "").strip())
                if best:
                    urls.append(abs_url(gw_url, best))
            else:
                src = (el.get("src") or "").strip()
                if src:
                    urls.append(abs_url(gw_url, src))

    urls = filter_gw_product_images(urls, keep_360=keep_360)
    urls = uniq_keep_order(urls)
    urls = dedupe_by_filename(urls)
    urls = keep_real_product_images(urls)

    urls = urls[:max_images]

    if ensure_query:
        urls = [ensure_query_defaults(u, ensure_query_default) for u in urls]

    return urls

def scrape_gw_images_fallback_regex(
    gw_url: str,
    html: str,
    max_images: int = 20,
    keep_360: bool = False,
    ensure_query: bool = False,
    ensure_query_default: str = "fm=webp&w=1200&h=1237",
) -> List[str]:
    urls: List[str] = []

    if not html:
        return urls

    patterns = [
        r'''srcset=["']([^"']*?/app/resources/catalog/product/[^"']+)["']''',
        r'''src=["']([^"']*?/app/resources/catalog/product/[^"']+)["']''',
        r'''srcset=["']([^"']*?/resources/catalog/product/[^"']+)["']''',
        r'''src=["']([^"']*?/resources/catalog/product/[^"']+)["']''',
        r'''srcset=["']([^"']*?/media/catalog/product/[^"']+)["']''',
        r'''src=["']([^"']*?/media/catalog/product/[^"']+)["']''',
        r'''srcset=["']([^"']*?/catalog/product/[^"']+)["']''',
        r'''src=["']([^"']*?/catalog/product/[^"']+)["']''',
    ]
    

    for pat in patterns:
        matches = re.findall(pat, html, flags=re.I)
        for match in matches:
            candidate = match.strip()

            if "," in candidate:
                parts = [x.strip() for x in candidate.split(",") if x.strip()]
                for part in parts:
                    url_part = part.split(" ")[0].strip()
                    if url_part:
                        urls.append(abs_url(gw_url, url_part))
            else:
                urls.append(abs_url(gw_url, candidate))

    urls = filter_gw_product_images(urls, keep_360=keep_360)
    urls = uniq_keep_order(urls)
    urls = dedupe_by_filename(urls)
    urls = keep_real_product_images(urls)

    if ensure_query:
        urls = [ensure_query_defaults(u, ensure_query_default) for u in urls]

    return urls[:max_images]

def keep_real_product_images(images: List[str]) -> List[str]:
    out: List[str] = []

    bad_markers = [
        "aeronautica_imperialis",
        "landscape",
        "header",
        "banner",
        "category",
    ]

    for u in images:
        low = u.lower()

        if any(marker in low for marker in bad_markers):
            continue

        looks_product_like = (
            "/catalog/product/" in low
            or "/app/resources/catalog/product/" in low
            or "/media/catalog/product/" in low
            or "/resources/catalog/product/" in low
        )

        if looks_product_like:
            out.append(u)

    return out if out else images


def extract_code_from_images(images: List[str]) -> str:
    """
    Pokusí se vytáhnout produktový kód z GW obrázků.
    Např. 99120206012_SkavenHellPitAbomination...
    """
    for u in images:
        m = re.search(r"/(\d{8,14})_[^/]+(?:\.jpg|\.jpeg|\.png|\.webp|\.avif)", u, flags=re.I)
        if m:
            return m.group(1)
    return ""


# =========================
# TEMPLATE RESOLUTION
# =========================
@dataclass
class LangTriple:
    cs: str
    en: str
    sk: str


@dataclass
class TemplatesMulti:
    short_default: LangTriple
    short_universal: LangTriple
    detail_mini: LangTriple
    detail_book: LangTriple
    detail_dice: LangTriple
    detail_warscroll: LangTriple
    detail_upgrades: LangTriple
    detail_accessories: LangTriple


def load_docx_if_exists(path: Path) -> Optional[str]:
    if path.exists():
        return load_docx_raw_text(str(path))
    return None


def load_lang_docx(base_dir: Path, stem: str, *, fallback_to_cs: bool = True) -> LangTriple:
    cs_path = base_dir / f"{stem}.docx"
    en_path = base_dir / f"{stem}_en.docx"
    sk_path = base_dir / f"{stem}_sk.docx"

    cs = load_docx_if_exists(cs_path)
    en = load_docx_if_exists(en_path)
    sk = load_docx_if_exists(sk_path)

    if cs is None:
        raise FileNotFoundError(f"Chybí povinná CZ šablona: {cs_path}")

    if fallback_to_cs:
        if en is None:
            en = cs
        if sk is None:
            sk = cs

    return LangTriple(
        cs=cs,
        en=en if en is not None else "",
        sk=sk if sk is not None else "",
    )


def load_templates_multi_from_dir(base_dir: str) -> TemplatesMulti:
    d = Path(base_dir).expanduser()

    return TemplatesMulti(
        short_default=load_lang_docx(d, "kratky"),
        short_universal=load_lang_docx(d, "kratky_univ"),
        detail_mini=load_lang_docx(d, "detailni"),
        detail_book=load_lang_docx(d, "detailni_kniha"),
        detail_dice=load_lang_docx(d, "detailni_kostky"),
        detail_warscroll=load_lang_docx(d, "detailni_warscoll"),
        detail_upgrades=load_lang_docx(d, "detailni_upgrades"),
        detail_accessories=load_lang_docx(d, "detailni_prislusentvi", fallback_to_cs=True)
        if (d / "detailni_prislusentvi.docx").exists()
        else load_lang_docx(d, "detailni"),
    )


def pick_short_template_by_type(tpl: TemplatesMulti, ptype: str) -> Tuple[str, LangTriple]:
    if ptype in {"book", "dice", "warscroll", "accessories", "upgrades"}:
        return "kratky_univ", tpl.short_universal
    return "kratky", tpl.short_default


def pick_detail_template_by_type(tpl: TemplatesMulti, ptype: str) -> Tuple[str, LangTriple]:
    if ptype == "upgrades":
        return "detailni_upgrades", tpl.detail_upgrades
    if ptype == "accessories":
        return (
            "detailni_prislusentvi"
            if tpl.detail_accessories != tpl.detail_mini
            else "detailni (fallback místo detailni_prislusentvi)",
            tpl.detail_accessories,
        )
    if ptype == "book":
        return "detailni_kniha", tpl.detail_book
    if ptype == "dice":
        return "detailni_kostky", tpl.detail_dice
    if ptype == "warscroll":
        return "detailni_warscoll", tpl.detail_warscroll
    return "detailni", tpl.detail_mini


# =========================
# SAVE SINGLE-ROW CSV
# =========================
def save_single_row_csv(row: Dict[str, str], columns: List[str], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    df_one = pd.DataFrame([row], columns=columns)
    df_one.to_csv(out_path, sep=";", index=False, encoding="utf-8-sig", quoting=csv.QUOTE_ALL)


# =========================
# RUN
# =========================
def run_scraper(
    input_links: str,
    output: str,
    tpl_dir: str,
    vat: float = 21.0,
    stock: str = "0",
    product_visibility: str = "hidden",
    availability_in_stock: str = "Skladem",
    availability_out_of_stock: str = "Na dotaz",
    oss_tax_rate_cz: str = "",
    percent_vat: str = "21",
    verbose: bool = False,
    warranty: str = "2 roky",
    supplier: str = "Games Workshop",
    google_category_id_in_feed: str = "1246",
    heureka_category_id: str = "1475",
    zbozi_category_id: str = "412",
    google_category_id: str = "1246",
    split_out_dir: str = "",
    split_name: str = "index_name",
    split_prefix: str = "",
    split_by_type: bool = True,
    filename_include_type: bool = False,
    max_images: int = 20,
    keep_360: bool = False,
    only_first_image: bool = False,
    images_ensure_query: bool = False,
    html_singleline: bool = False,
) -> Dict[str, object]:
    links = load_links(input_links)
    tpl = load_templates_multi_from_dir(tpl_dir)

    create_rows: List[Dict[str, str]] = []
    source_rows: List[Dict[str, str]] = []
    processed_files: List[str] = []
    debug_logs: List[str] = []

    split_dir = Path(split_out_dir).expanduser() if split_out_dir else None
    if split_dir:
        split_dir.mkdir(parents=True, exist_ok=True)

    total = len(links)

    for i, r in links.iterrows():
        idx = i + 1
        hp_url = str(r["hp_url"]).strip()
        gw_url = str(r["gw_url"]).strip()

        hp_html, _hp_final = fetch_html(hp_url, referer="https://www.herniprostor.cz/", headers=BASE_HEADERS)
        hp_soup = BeautifulSoup(hp_html, "html.parser")

        h1_raw = hp_extract_h1(hp_soup)
        h1 = clean_title(h1_raw)
        _system = hp_pick_system(hp_soup)

        faction = ""
        name_final = build_name_from_h1(h1, _system)

        price = hp_extract_price(hp_soup) or 0.0
        ean = hp_extract_ean(hp_soup)
        external = hp_extract_external_code(hp_soup)

        categoryText = "Warhammer"
        ptype = detect_product_type(name_final, categoryText, ean)

        short_tpl_name, short_tpl = pick_short_template_by_type(tpl, ptype)
        detail_tpl_name, detail_tpl = pick_detail_template_by_type(tpl, ptype)

        short_cs = short_tpl.cs
        short_en = short_tpl.en
        short_sk = short_tpl.sk

        detail_cs = detail_tpl.cs
        detail_en = detail_tpl.en
        detail_sk = detail_tpl.sk

        if html_singleline:
            short_cs = safe_singleline_html(short_cs)
            short_en = safe_singleline_html(short_en)
            short_sk = safe_singleline_html(short_sk)
            detail_cs = safe_singleline_html(detail_cs)
            detail_en = safe_singleline_html(detail_en)
            detail_sk = safe_singleline_html(detail_sk)

        images: List[str] = []
        gw_price: Optional[float] = None
        gw_currency: Optional[str] = None
        gw_final = gw_url

        if gw_url:
            try:
                gw_html, gw_final = fetch_gw_html(gw_url, verbose=verbose)
                gw_price, gw_currency = gw_extract_price(gw_html)
                faction = detect_faction_from_gw_html(gw_html, _system)

                images = scrape_gw_images_stable(
                    gw_final,
                    gw_html,
                    max_images=max_images,
                    keep_360=bool(keep_360),
                    ensure_query=bool(images_ensure_query),
                )
                if verbose:
                    print("----- IMAGE SCRAPE DEBUG -----")
                    print("stable images:", len(images))
                    for u in images[:10]:
                        print(" stable:", u)
                    print("------------------------------")

                if not images:
                    log("GW stable image scrape returned 0 images, trying regex fallback...", verbose)
                    images = scrape_gw_images_fallback_regex(
                        gw_final,
                        gw_html,
                        max_images=max_images,
                        keep_360=bool(keep_360),
                        ensure_query=bool(images_ensure_query),
                    )

                    if verbose:
                        print("----- REGEX FALLBACK DEBUG -----")
                        print("regex images:", len(images))
                        for u in images[:10]:
                            print(" regex:", u)
                        print("--------------------------------")

                images = keep_real_product_images(images)

                if verbose:
                    print("----- AFTER keep_real_product_images -----")
                    print("final images:", len(images))
                    for u in images[:10]:
                        print(" final:", u)
                    print("------------------------------------------")

            except Exception as e:
                log(f"GW ERROR: {e}", verbose)
                images = []
                gw_price = None
                gw_currency = None
                faction = ""

        if only_first_image and images:
            images = [images[0]]

        fallback_code = extract_code_from_images(images)

        if not is_valid_product_code(external):
            external = ""

        if not external and fallback_code:
            external = fallback_code

        if ean:
            code = ean
        elif is_valid_product_code(external):
            code = external
        elif fallback_code:
            code = fallback_code
        else:
            code = generate_fallback_code_from_url(hp_url, gw_final if gw_url else "")

        name_final = build_name_from_h1(h1, _system, faction)

        xml_feed_name = make_xml_feed_name_from_name(name_final)
        seo_title = make_seo_title_from_name(name_final)

        std_price = round(gw_price * FX_RATES[gw_currency], 2) if (gw_price is not None and gw_currency in FX_RATES) else None

        if verbose:
            debug_block = f"""----- KONTROLA DAT -----
Raw název z GW: {h1}
Finální název: {name_final}
Systém: {_system}
Frakce: {faction or '-'}
Typ produktu: {ptype}
Code: {code or '-'}
Standardní cena: {fmt_cz_money(std_price)} Kč
Prodejní cena: {fmt_cz_money(price)} Kč
Počet obrázků: {len(images)}
------------------------
"""
            print(debug_block)
            debug_logs.append(debug_block)

        print(f"[{idx:02d}/{total:02d}] {name_final}")
        print(f"  HP: {hp_url}")
        print(f"  GW: {gw_final if gw_url else '-'}")
        print(
            f"  type={ptype}"
            f" | code={(code or '-')}"
            f" | ean={(ean or '-')}"
            f" | external={(external or '-')}"
            f" | price={fmt_cz_money(price)}"
            f" | gwPrice={(f'{gw_price}{gw_currency}' if gw_price is not None and gw_currency else '-')}"
            f" | stdPrice={(fmt_cz_money(std_price) if std_price is not None else '-')}"
            f" | imgs={len(images)}"
        )
        if images:
            print(f"  first image: {images[0]}")
        else:
            print("  first image: -")
        print(f"  templates: short={short_tpl_name} | detail={detail_tpl_name}")
        print("  ✅ OK\n")

        create_row: Dict[str, str] = {c: "" for c in CREATE_COLUMNS}
        create_row.update({
            "code": code,
            "pairCode": "",
            "name": name_final,
            "price": f"{price:.2f}".replace(".", ","),
            "description": "",
        })

        source_row: Dict[str, str] = {c: "" for c in SOURCE_COLUMNS}
        source_row.update({
            "code": code,
            "pairCode": "",
            "name": name_final,
            "ean": ean,
            "externalCode": external,
            "price": f"{price:.2f}".replace(".", ","),
            "priceWithoutVat": f"{price_without_vat(price, float(vat)):.2f}".replace(".", ","),
            "standardPrice": f"{std_price:.2f}".replace(".", ",") if std_price is not None else "",
            "categoryText": categoryText,
            "warranty": str(warranty),
            "supplier": str(supplier),
            "googleCategoryIdInFeed": str(google_category_id_in_feed),
            "heurekaCategoryId": str(heureka_category_id),
            "zboziCategoryId": str(zbozi_category_id),
            "googleCategoryId": str(google_category_id),
            "stock": str(stock),
            "percentVat": str(percent_vat),
            "ossTaxRate:CZ": str(oss_tax_rate_cz),
            "availabilityInStock": str(availability_in_stock),
            "availabilityOutOfStock": str(availability_out_of_stock),
            "productVisibility": str(product_visibility),
            "xmlFeedName": xml_feed_name,
            "seoTitle": seo_title,
            "system": _system,
            "faction": faction,
            "productType": ptype,
            "hp_url": hp_url,
            "gw_url": gw_final if gw_url else "",
        })

        if images:
            create_row["image"] = images[0]
            source_row["image"] = images[0]

            for idx_img in range(1, min(20, len(images))):
                col = f"image{idx_img + 1}"
                create_row[col] = images[idx_img]
                source_row[col] = images[idx_img]

        create_rows.append(create_row)
        source_rows.append(source_row)

        if split_dir:
            safe_name = slugify_filename(name_final)
            safe_ean = slugify_filename(ean) if ean else "bez_ean"
            safe_code = slugify_filename(code) if code else "bez_kodu"

            if split_name == "name":
                base_name = f"{safe_name}_{safe_ean}"
            elif split_name == "code":
                base_name = safe_code
            elif split_name == "index_code":
                base_name = f"{idx:02d}_{safe_code}"
            else:
                base_name = f"{idx:02d}_{safe_name}_{safe_ean}"

            if split_prefix:
                base_name = f"{slugify_filename(split_prefix)}{base_name}"

            if filename_include_type:
                base_name = f"{ptype}__{base_name}"

            out_folder = split_dir
            if split_by_type:
                out_folder = split_dir / ptype
                out_folder.mkdir(parents=True, exist_ok=True)

            create_out_path = out_folder / f"{base_name}_CREATE.csv"
            source_out_path = out_folder / f"{base_name}_SOURCE.csv"

            save_single_row_csv(create_row, CREATE_COLUMNS, create_out_path)
            save_single_row_csv(source_row, SOURCE_COLUMNS, source_out_path)

            processed_files.append(str(create_out_path))
            processed_files.append(str(source_out_path))

    output_path = Path(output).expanduser()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    create_output_path = output_path.with_name(f"{output_path.stem}_CREATE{output_path.suffix}")
    source_output_path = output_path.with_name(f"{output_path.stem}_SOURCE{output_path.suffix}")

    create_df = pd.DataFrame(create_rows, columns=CREATE_COLUMNS)
    create_df = create_df.reindex(columns=CREATE_COLUMNS, fill_value="")
    create_df.to_csv(create_output_path, sep=";", index=False, encoding="utf-8-sig", quoting=csv.QUOTE_ALL)

    source_df = pd.DataFrame(source_rows, columns=SOURCE_COLUMNS)
    source_df = source_df.reindex(columns=SOURCE_COLUMNS, fill_value="")
    source_df.to_csv(source_output_path, sep=";", index=False, encoding="utf-8-sig", quoting=csv.QUOTE_ALL)

    print(f"\nOK: saved CREATE {create_output_path} ({len(create_df)} rows)")
    print(f"OK: saved SOURCE {source_output_path} ({len(source_df)} rows)")
    if split_dir:
        print(f"OK: saved split CSVs to {split_dir} ({len(processed_files)} files)")

    return {
        "create_output_csv": str(create_output_path),
        "source_output_csv": str(source_output_path),
        "split_dir": str(split_dir) if split_dir else "",
        "row_count": len(create_df),
        "create_rows": create_rows,
        "source_rows": source_rows,
        "split_files": processed_files,
        "debug_logs": debug_logs,
    }